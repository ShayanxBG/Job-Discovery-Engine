#!/usr/bin/env python3
import json, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BLUE = '1F4E79'
GREY = 'BFBFBF'


def set_bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), GREY)
    pbdr.append(bot)
    pPr.append(pbdr)


def set_run_font(run, size=None, bold=None, italic=None, colour=None):
    run.font.name = 'Calibri'
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if colour:
        run.font.color.rgb = RGBColor.from_string(colour)


def compact(p, before=0, after=0, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    return p


def split_header(header):
    parts = [x.strip() for x in str(header).split('|')]
    if len(parts) >= 2:
        return parts[0], ' | '.join(parts[1:])
    return str(header), ''


def add_entry_header(d, header, dates):
    role, employer = split_header(header)
    p = compact(d.add_paragraph(), before=2.5, after=1.5, line=1.0)
    set_run_font(p.add_run(role), size=9.6, bold=True)
    if employer:
        set_run_font(p.add_run(' | ' + employer), size=9.6)
    if dates:
        set_run_font(p.add_run(' | ' + str(dates)), size=9.35)
    return p


def add_section_heading(d, title):
    p = compact(d.add_paragraph(), before=3.0, after=1.5, line=1.0)
    r = p.add_run(title)
    set_run_font(r, size=10.6, bold=True, colour=BLUE)
    set_bottom_border(p)
    return p


def add_bullet(d, text):
    p = compact(d.add_paragraph(style='List Bullet'), before=0, after=0.3, line=1.0)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.08)
    set_run_font(p.add_run(text), size=9.25)
    return p


def main():
    if len(sys.argv) != 3:
        raise SystemExit('Usage: python tools/render_cv_docx.py input.json output.docx')

    data = json.loads(Path(sys.argv[1]).read_text(encoding='utf-8'))
    out = Path(sys.argv[2])
    out.parent.mkdir(parents=True, exist_ok=True)

    d = Document()
    sec = d.sections[0]
    sec.top_margin = Inches(0.38)
    sec.bottom_margin = Inches(0.38)
    sec.left_margin = Inches(0.625)
    sec.right_margin = Inches(0.625)

    normal = d.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(9.4)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # Make Word's built-in List Bullet compact too.
    lb = d.styles['List Bullet']
    lb.font.name = 'Calibri'
    lb.font.size = Pt(9.25)
    lb.paragraph_format.space_before = Pt(0)
    lb.paragraph_format.space_after = Pt(0)
    lb.paragraph_format.line_spacing = 1.0

    p = compact(d.add_paragraph(), after=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(data['name']), size=17, bold=True)

    p = compact(d.add_paragraph(), after=2.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data['contact'])
    set_run_font(r, size=9.1, colour='0563C1')

    add_section_heading(d, 'PROFESSIONAL SUMMARY')
    p = compact(d.add_paragraph(), after=0.5, line=1.0)
    set_run_font(p.add_run(data['summary']), size=9.45)

    for secdata in data.get('sections', []):
        add_section_heading(d, secdata['title'])
        for e in secdata.get('entries', []):
            add_entry_header(d, e.get('header', ''), e.get('dates', ''))
            for line in e.get('lines', []):
                p = compact(d.add_paragraph(), after=0.5, line=1.0)
                set_run_font(p.add_run(line), size=9.15, italic=True)
            for b in e.get('bullets', []):
                add_bullet(d, b)

        for line in secdata.get('lines', []):
            p = compact(d.add_paragraph(), after=0.15, line=1.0)
            if ':' in line:
                label, rest = line.split(':', 1)
                set_run_font(p.add_run(label + ':'), size=9.15, bold=True)
                set_run_font(p.add_run(rest), size=9.15)
            else:
                set_run_font(p.add_run(line), size=9.15)

    d.save(out)
    print(out)


if __name__ == '__main__':
    main()
